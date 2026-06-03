import re
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cache for optimized East Asian font setup
_east_asian_font_cache = {}


def _insert_latex_to_paragraph(p, latex_content, is_block=False):
    """Insert a LaTeX formula into a docx paragraph as styled text"""
    clean = latex_content.strip()
    if is_block:
        # 块级公式：独立成段，居中显示
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clean)
        run.font.name = 'Cambria Math'
        run.font.size = docx.shared.Pt(14)
        run.font.color.rgb = docx.shared.RGBColor(0, 51, 102)
    else:
        # 行内公式
        run = p.add_run(f" {clean} ")
        run.font.name = 'Cambria Math'
        run.font.size = docx.shared.Pt(11)


def _insert_latex_to_heading(p, latex_content):
    """Insert a LaTeX formula into a docx heading paragraph as styled text"""
    clean = latex_content.strip()
    run = p.add_run(f" {clean} ")
    run.font.name = 'Cambria Math'


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph run"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # This gets access to the document.xml.rels file
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add it to the paragraph
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object
    new_run = docx.text.run.Run(
        docx.oxml.OxmlElement('w:r'),
        paragraph._element
    )
    new_run.text = text
    new_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    new_run.font.underline = True

    # Add the run to the hyperlink
    hyperlink.append(new_run._element)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)

    return new_run


def add_latex_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with LaTeX formula support - optimized version"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = docx.shared.Pt(0)
    p.paragraph_format.space_after = docx.shared.Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Split by $$...$$ (block) and $...$ (inline)
    parts = re.split(r'(\$\$.*?\$\$|\$[^$\n]+?\$)', text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue
        if part.startswith('$$') and part.endswith('$$'):
            latex_content = part[2:-2].strip()
            _insert_latex_to_paragraph(p, latex_content, is_block=True)
        elif part.startswith('$') and part.endswith('$') and len(part) > 2:
            latex_content = part[1:-1].strip()
            _insert_latex_to_paragraph(p, latex_content, is_block=False)
        else:
            # Optimized inline formatting with single-pass processing
            # Use regex to find all formatting markers at once
            formatting_pattern = r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`|https?://[^\s<>"\')\]]+)'
            segments = re.split(formatting_pattern, part, flags=re.DOTALL)

            # Batch process segments to minimize XML operations
            runs_to_add = []
            for segment in segments:
                if not segment:
                    continue

                if segment.startswith('**') and segment.endswith('**') and len(segment) > 4:
                    content = segment[2:-2]
                    runs_to_add.append(('bold', content))
                elif segment.startswith('*') and segment.endswith('*') and len(segment) > 2:
                    content = segment[1:-1]
                    runs_to_add.append(('italic', content))
                elif segment.startswith('`') and segment.endswith('`') and len(segment) > 2:
                    content = segment[1:-1]
                    runs_to_add.append(('code', content))
                elif re.match(r'https?://', segment):
                    url = segment.rstrip('.,;')
                    runs_to_add.append(('url', url))
                else:
                    runs_to_add.append(('text', segment))

            # Add all runs in batch
            for run_type, content in runs_to_add:
                if run_type == 'bold':
                    run = p.add_run(content)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    _set_east_asian_font(run, '宋体')
                elif run_type == 'italic':
                    run = p.add_run(content)
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    _set_east_asian_font(run, '宋体')
                elif run_type == 'code':
                    run = p.add_run(content)
                    run.font.name = 'Courier New'
                    run.font.size = docx.shared.Pt(10)
                    _set_east_asian_font(run, '宋体')
                elif run_type == 'url':
                    add_hyperlink(p, content, content)
                else:  # text
                    run = p.add_run(content)
                    run.font.name = 'Times New Roman'
                    _set_east_asian_font(run, '宋体')

    return p


def add_latex_heading(doc, text, level=1):
    """Add a heading that may contain LaTeX formulas - optimized version"""
    heading = doc.add_heading('', level=level)
    heading.paragraph_format.space_before = docx.shared.Pt(12)
    heading.paragraph_format.space_after = docx.shared.Pt(6)

    # Configure font properties based on heading level
    font_size = {1: docx.shared.Pt(16), 2: docx.shared.Pt(14), 3: docx.shared.Pt(12)}.get(level, docx.shared.Pt(12))
    east_asian_font = '黑体'

    parts = re.split(r'(\$\$.*?\$\$|\$[^$\n]+?\$)', text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue
        if part.startswith('$$') and part.endswith('$$'):
            latex_content = part[2:-2].strip()
            _insert_latex_to_heading(heading, latex_content)
        elif part.startswith('$') and part.endswith('$') and len(part) > 2:
            latex_content = part[1:-1].strip()
            _insert_latex_to_heading(heading, latex_content)
        else:
            run = heading.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = font_size
            run.bold = True
            _set_east_asian_font(run, east_asian_font)

    return heading


def _set_east_asian_font(font_obj, font_name):
    """Safely set the East Asian font for a docx Font object with caching"""
    # Use font object id as cache key to avoid repeated XML manipulation
    font_id = id(font_obj)
    if font_id in _east_asian_font_cache and _east_asian_font_cache[font_id] == font_name:
        return

    rpr = font_obj._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = docx.oxml.OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)

    # Cache the result
    _east_asian_font_cache[font_id] = font_name
