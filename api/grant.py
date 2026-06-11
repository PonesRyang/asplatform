from __future__ import annotations

import json
import io
import re
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from api.dependencies import verify_service_access
from database import get_db
from models import AdminUser, GrantConfigItem, GrantProject, GrantStep
from schemas.grant import (
    GrantConfigOption,
    GrantConfigOptionsResponse,
    GrantConfigTreeNode,
    GrantInputState,
    GrantKeywordPatch,
    GrantProjectCreate,
    GrantProjectResponse,
    GrantProjectSummary,
    GrantProjectUpdate,
    GrantStepAction,
    GrantStepHistoryItem,
    GrantTopicSelect,
)
from utils.auth import check_permission, get_optional_admin, verify_token
from services import ai_service, literature_service

router = APIRouter(prefix="/api/ai/grant", tags=["grant"])


def _json_dumps(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False)


def _json_loads(value: Optional[str], default: Any):
    if not value:
        return default
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return default


def _input_to_dict(item: GrantInputState) -> Dict[str, Any]:
    return {
        "fundType": item.fund_type,
        "researchAreaPath": item.research_area_path,
        "subject": item.subject or "",
        "diseasePath": item.disease_path or [],
        "phenotype": item.phenotype or "",
        "variableType": item.variable_type or "",
        "variableName": item.variable_name or "",
    }


def _project_input(project: GrantProject) -> Dict[str, Any]:
    return {
        "fundType": project.fund_type,
        "researchAreaPath": _json_loads(project.research_area_path, []),
        "subject": project.subject or "",
        "diseasePath": _json_loads(project.disease_path, []),
        "phenotype": project.phenotype or "",
        "variableType": project.variable_type or "",
        "variableName": project.variable_name or "",
    }


def _project_to_response(project: GrantProject) -> GrantProjectResponse:
    return GrantProjectResponse(
        id=project.id,
        title=project.title,
        status=project.status,
        current_step=project.current_step,
        input=_project_input(project),
        keywords=_json_loads(project.keywords_json, {"must": [], "should": [], "groups": []}),
        references=_json_loads(project.references_json, []),
        topics=_json_loads(project.topics_json, []),
        report_sections=_json_loads(project.report_sections_json, []),
        proposal_sections=_json_loads(project.proposal_sections_json, []),
        created_at=project.created_at,
        updated_at=project.updated_at,
    )


def _project_to_summary(project: GrantProject) -> GrantProjectSummary:
    return GrantProjectSummary(
        id=project.id,
        title=project.title,
        status=project.status,
        current_step=project.current_step,
        fund_type=project.fund_type,
        research_area_path=_json_loads(project.research_area_path, []),
        updated_at=project.updated_at,
    )


def _flat_config_options(db: Session, category: str) -> List[GrantConfigOption]:
    items = db.query(GrantConfigItem).filter(
        GrantConfigItem.category == category,
        GrantConfigItem.is_active == True,
        GrantConfigItem.parent_id.is_(None),
    ).order_by(GrantConfigItem.sort_order.asc(), GrantConfigItem.id.asc()).all()
    return [GrantConfigOption(label=item.label, value=item.value or item.label) for item in items]


def _tree_config_options(db: Session, category: str) -> List[GrantConfigTreeNode]:
    items = db.query(GrantConfigItem).filter(
        GrantConfigItem.category == category,
        GrantConfigItem.is_active == True,
    ).order_by(GrantConfigItem.sort_order.asc(), GrantConfigItem.id.asc()).all()

    node_by_id: Dict[int, GrantConfigTreeNode] = {
        item.id: GrantConfigTreeNode(label=item.label, value=item.value or item.label)
        for item in items
    }
    roots: List[GrantConfigTreeNode] = []

    for item in items:
        node = node_by_id[item.id]
        if item.parent_id and item.parent_id in node_by_id:
            node_by_id[item.parent_id].children.append(node)
        else:
            roots.append(node)

    return roots


async def _authorize_project_query(
    db: Session,
    token: Optional[str],
    current_user: Optional[AdminUser],
):
    if current_user:
        check_permission(current_user, "ai")
        query = db.query(GrantProject)
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(GrantProject.admin_id == current_user.id)
        return query, None

    if not token:
        raise HTTPException(status_code=401, detail="Access token required")
    token_record = verify_token(token, db, "ai")
    return db.query(GrantProject).filter(GrantProject.token_id == token_record.id), token_record


async def _get_project(
    db: Session,
    project_id: int,
    token: Optional[str],
    current_user: Optional[AdminUser],
) -> GrantProject:
    query, _ = await _authorize_project_query(db, token, current_user)
    project = query.filter(GrantProject.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Grant project not found")
    return project


def _save_step(db: Session, project: GrantProject, step_key: str, output: Any, status: str = "ready"):
    step = GrantStep(
        project_id=project.id,
        step_key=step_key,
        status=status,
        input_json=_json_dumps(_project_input(project)),
        output_json=_json_dumps(output),
    )
    db.add(step)
    project.current_step = step_key
    project.updated_at = datetime.now(timezone.utc)


def _strip_json_fence(content: str) -> str:
    text = content.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


async def _require_ai_json(prompt: str) -> Any:
    response = await ai_service.chat_completion([{"role": "user", "content": prompt}], temperature=0.5)
    content = response.get("content", "")
    if not content or content.startswith("[Mock Response]") or content.startswith("AI "):
        raise HTTPException(status_code=502, detail=content or "AI 生成失败，未写入生成内容")

    try:
        return json.loads(_strip_json_fence(content))
    except Exception:
        raise HTTPException(status_code=502, detail="AI 返回内容不是合法 JSON，未写入生成内容，请重新生成")


def _reference_search_query(project: GrantProject) -> str:
    keywords = _json_loads(project.keywords_json, {"must": [], "should": []})
    terms: List[str] = []

    for keyword in keywords.get("must", []):
        if keyword.get("selected", True) and keyword.get("text"):
            terms.append(str(keyword["text"]))

    for keyword in keywords.get("should", []):
        if keyword.get("selected", True) and keyword.get("text"):
            terms.append(str(keyword["text"]))

    for value in [
        project.subject,
        project.phenotype,
        project.variable_name,
        *(_json_loads(project.disease_path, []) or []),
    ]:
        if value:
            terms.append(str(value))

    deduped: List[str] = []
    seen = set()
    for term in terms:
        normalized = term.strip()
        if normalized and normalized.lower() not in seen:
            deduped.append(normalized)
            seen.add(normalized.lower())

    return " ".join(deduped[:8])


def _citation_year(value: Any) -> Optional[int]:
    if value is None:
        return None
    text = json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value)
    match = re.search(r"\b(19|20)\d{2}\b", text)
    if match:
        return int(match.group(0))
    return None


def _citation_to_reference(citation: Dict[str, Any], index: int, query: str) -> Optional[Dict[str, Any]]:
    title = str(citation.get("title") or "").strip()
    if not title:
        return None

    link = citation.get("link") or ""
    pmid = citation.get("pmid") or citation.get("PMID") or citation.get("id") or ""
    if not pmid and "pubmed.ncbi.nlm.nih.gov/" in link:
        pmid = link.rstrip("/").split("/")[-1]

    journal = citation.get("source") or citation.get("journal") or citation.get("container-title") or ""
    if isinstance(journal, list):
        journal = journal[0] if journal else ""

    database = citation.get("database") or "学术数据库"
    year = _citation_year(citation.get("year") or citation.get("published") or citation.get("pubdate"))

    return {
        "id": f"ref-{index}",
        "pmid": str(pmid or ""),
        "doi": str(citation.get("doi") or citation.get("DOI") or ""),
        "title": title,
        "journal": str(journal or ""),
        "year": year,
        "evidenceNote": f"来自{database}，与检索式「{query}」相关，可用于支撑选题依据与研究背景。",
        "selectedForGeneration": index <= 5,
        "link": str(link or ""),
        "database": str(database),
        "formatted": str(citation.get("formatted") or ""),
    }


async def _search_real_references(project: GrantProject) -> List[Dict[str, Any]]:
    query = _reference_search_query(project)
    if not query:
        return []

    try:
        citations = await literature_service.search_literature(query, max_results=10)
    except Exception:
        return []

    references: List[Dict[str, Any]] = []
    seen_titles = set()
    for citation in citations:
        reference = _citation_to_reference(citation, len(references) + 1, query)
        if not reference:
            continue
        normalized_title = reference["title"].lower()
        if normalized_title in seen_titles:
            continue
        seen_titles.add(normalized_title)
        references.append(reference)

    return references


def _grant_context(project: GrantProject) -> str:
    return "\n".join([
        f"课题类型：{project.fund_type}",
        f"研究方向：{' / '.join(_json_loads(project.research_area_path, []))}",
        f"申请书主题：{project.subject or ''}",
        f"疾病：{' / '.join(_json_loads(project.disease_path, []))}",
        f"表型/科学问题：{project.phenotype or ''}",
        f"主变量：{project.variable_type or ''} / {project.variable_name or ''}",
    ])


def _add_docx_paragraph(document, text: str, style: Optional[str] = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    return paragraph


def _build_proposal_docx(project: GrantProject) -> io.BytesIO:
    document = docx.Document()
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    title = document.add_heading(project.title or "基金申请书", level=0)
    title.alignment = 1

    _add_docx_paragraph(document, f"项目类型：{project.fund_type or ''}")
    _add_docx_paragraph(document, f"研究方向：{' / '.join(_json_loads(project.research_area_path, []))}")
    _add_docx_paragraph(document, f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    document.add_paragraph("")
    sections = _json_loads(project.proposal_sections_json, [])
    if not sections:
        _add_docx_paragraph(document, "暂无申请书正文，请先生成基金申请书。")
    for section in sections:
        document.add_heading(section.get("title", "未命名章节"), level=1)
        markdown = section.get("markdown", "")
        for paragraph in [line.strip() for line in markdown.split("\n") if line.strip()]:
            _add_docx_paragraph(document, paragraph)

    references = _json_loads(project.references_json, [])
    if references:
        document.add_heading("参考文献", level=1)
        for index, ref in enumerate(references, start=1):
            _add_docx_paragraph(
                document,
                f"{index}. {ref.get('title', '')}. {ref.get('journal', '')}, {ref.get('year', '')}. DOI: {ref.get('doi', '')}",
            )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("/config/options", response_model=GrantConfigOptionsResponse)
async def get_grant_config_options(
    token: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    if current_user:
        check_permission(current_user, "ai")
    else:
        await verify_service_access(db, token, "ai")

    return GrantConfigOptionsResponse(
        fundTypes=_flat_config_options(db, "fund_type"),
        researchAreas=_tree_config_options(db, "research_area"),
        diseases=_tree_config_options(db, "disease"),
        variableTypes=_flat_config_options(db, "variable_type"),
        phenotypes=_flat_config_options(db, "phenotype"),
    )


@router.post("/projects", response_model=GrantProjectResponse)
async def create_grant_project(
    item: GrantProjectCreate,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    token_record = None
    if current_user:
        check_permission(current_user, "ai")
    else:
        token_record = await verify_service_access(db, item.token, "ai")

    input_data = item.input
    title = f"{input_data.variable_name or input_data.variable_type or '核心变量'}相关{input_data.phenotype or '科学问题'}课题申报"
    if input_data.subject:
        title = input_data.subject.strip()

    project = GrantProject(
        token_id=token_record.id if token_record else None,
        admin_id=current_user.id if current_user else None,
        title=title,
        status="draft",
        current_step="input",
        fund_type=input_data.fund_type,
        research_area_path=_json_dumps(input_data.research_area_path),
        subject=input_data.subject,
        disease_path=_json_dumps(input_data.disease_path),
        phenotype=input_data.phenotype,
        variable_type=input_data.variable_type,
        variable_name=input_data.variable_name,
        keywords_json=_json_dumps({"must": [], "should": [], "groups": []}),
        references_json=_json_dumps([]),
        topics_json=_json_dumps([]),
        report_sections_json=_json_dumps([]),
        proposal_sections_json=_json_dumps([]),
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.get("/projects", response_model=List[GrantProjectSummary])
async def list_grant_projects(
    token: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    query, _ = await _authorize_project_query(db, token, current_user)
    projects = query.order_by(GrantProject.updated_at.desc()).all()
    return [_project_to_summary(project) for project in projects]


@router.get("/projects/{project_id}", response_model=GrantProjectResponse)
async def get_grant_project(
    project_id: int,
    token: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, token, current_user)
    return _project_to_response(project)


@router.patch("/projects/{project_id}", response_model=GrantProjectResponse)
async def update_grant_project(
    project_id: int,
    item: GrantProjectUpdate,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    if item.title is not None:
        project.title = item.title
    if item.current_step is not None:
        project.current_step = item.current_step
    if item.status is not None:
        project.status = item.status
    if item.input is not None:
        input_data = item.input
        project.fund_type = input_data.fund_type
        project.research_area_path = _json_dumps(input_data.research_area_path)
        project.subject = input_data.subject
        project.disease_path = _json_dumps(input_data.disease_path)
        project.phenotype = input_data.phenotype
        project.variable_type = input_data.variable_type
        project.variable_name = input_data.variable_name
    project.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/keywords/generate", response_model=GrantProjectResponse)
async def generate_keywords(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    prompt = f"""你是基金申报选题助手。请基于以下申报信息生成关键词结构，直接输出 JSON，不要输出 Markdown。

{_grant_context(project)}

JSON 格式：
{{
  "must": [{{"id": "must-1", "text": "必须包含关键词", "source": "ai", "selected": true}}],
  "should": [{{"id": "or-1", "text": "可选关键词", "source": "ai", "selected": true, "groupKey": "disease"}}],
  "groups": [{{"key": "disease", "label": "关联疾病", "keywords": [{{"id": "g1", "text": "关键词", "source": "ai", "selected": true}}]}}]
}}
"""
    keywords = await _require_ai_json(prompt)
    if not isinstance(keywords, dict) or not {"must", "should", "groups"}.issubset(keywords.keys()):
        raise HTTPException(status_code=502, detail="AI 返回的关键词结构不完整，未写入生成内容")
    project.keywords_json = _json_dumps(keywords)
    project.status = "keywords_ready"
    _save_step(db, project, "keywords", keywords)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.patch("/projects/{project_id}/keywords", response_model=GrantProjectResponse)
async def update_keywords(
    project_id: int,
    item: GrantKeywordPatch,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    project.keywords_json = _json_dumps(item.keywords)
    project.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/references/search", response_model=GrantProjectResponse)
async def search_references(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    references = await _search_real_references(project)
    project.references_json = _json_dumps(references)
    project.status = "references_ready"
    _save_step(db, project, "references", references)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/topics/generate", response_model=GrantProjectResponse)
async def generate_topics(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    prompt = f"""你是国家自然科学基金选题顾问。请基于申报信息、关键词和参考文献生成 10 个候选选题，直接输出 JSON 数组，不要输出 Markdown。

{_grant_context(project)}

关键词：
{project.keywords_json or ""}

参考文献：
{project.references_json or ""}

每个数组元素格式：
{{
  "id": "topic-1",
  "title": "题目",
  "description": "一句话说明",
  "innovation": "创新点",
  "feasibility": "可行性",
  "fundFit": "基金匹配度",
  "risk": "风险",
  "score": {{"innovation": 85, "feasibility": 80, "fundFit": 88, "evidence": 78}},
  "referenceIds": ["ref-1"],
  "selected": false
}}
"""
    topics = await _require_ai_json(prompt)
    if isinstance(topics, dict) and isinstance(topics.get("topics"), list):
        topics = topics["topics"]
    if not isinstance(topics, list) or not topics:
        raise HTTPException(status_code=502, detail="AI 返回的候选题结构不完整，未写入生成内容")
    if not any(topic.get("selected") for topic in topics if isinstance(topic, dict)):
        topics[0]["selected"] = True
    project.topics_json = _json_dumps(topics)
    project.status = "topics_ready"
    project.title = topics[0]["title"]
    _save_step(db, project, "topics", topics)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/topics/{topic_id}/select", response_model=GrantProjectResponse)
async def select_topic(
    project_id: int,
    topic_id: str,
    item: GrantTopicSelect,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    topics = _json_loads(project.topics_json, [])
    found = False
    for topic in topics:
        topic["selected"] = topic.get("id") == topic_id
        if topic["selected"]:
            project.title = topic.get("title", project.title)
            found = True
    if not found:
        raise HTTPException(status_code=404, detail="Topic not found")
    project.topics_json = _json_dumps(topics)
    project.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/report/generate", response_model=GrantProjectResponse)
async def generate_report(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    prompt = f"""你是基金选题报告专家。请为已选题目生成选题报告章节，直接输出 JSON 数组，不要输出 Markdown。

{_grant_context(project)}

候选题：
{project.topics_json or ""}

参考文献：
{project.references_json or ""}

数组元素格式：
{{"key": "purpose", "title": "研究目的、意义", "markdown": "正文"}}
必须包含：研究目的、意义；研究内容及实现方案；科学问题和科学假说；选题评估。
"""
    sections = await _require_ai_json(prompt)
    if isinstance(sections, dict) and isinstance(sections.get("sections"), list):
        sections = sections["sections"]
    if not isinstance(sections, list) or not sections:
        raise HTTPException(status_code=502, detail="AI 返回的选题报告结构不完整，未写入生成内容")
    project.report_sections_json = _json_dumps(sections)
    project.status = "report_ready"
    _save_step(db, project, "report", sections)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.get("/projects/{project_id}/report/history", response_model=List[GrantStepHistoryItem])
async def get_report_history(
    project_id: int,
    token: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, token, current_user)
    steps = (
        db.query(GrantStep)
        .filter(GrantStep.project_id == project.id, GrantStep.step_key == "report")
        .order_by(GrantStep.created_at.desc(), GrantStep.id.desc())
        .all()
    )
    return [
        GrantStepHistoryItem(
            id=step.id,
            step_key=step.step_key,
            status=step.status,
            output=_json_loads(step.output_json, []),
            created_at=step.created_at,
        )
        for step in steps
    ]


@router.post("/projects/{project_id}/proposal/generate", response_model=GrantProjectResponse)
async def generate_proposal(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    prompt = f"""你是国家自然科学基金申请书写作专家。请基于选题报告生成申请书初稿章节，直接输出 JSON 数组，不要输出 Markdown。

{_grant_context(project)}

选题报告：
{project.report_sections_json or ""}

数组元素格式：
{{"key": "abstract", "title": "中文摘要、关键词", "status": "ready", "wordCount": 500, "markdown": "正文"}}
必须包含中文摘要、科学问题属性选择理由、项目立项依据、研究内容、技术路线图、年度研究计划及预期结果。
"""
    sections = await _require_ai_json(prompt)
    if isinstance(sections, dict) and isinstance(sections.get("sections"), list):
        sections = sections["sections"]
    if not isinstance(sections, list) or not sections:
        raise HTTPException(status_code=502, detail="AI 返回的申请书结构不完整，未写入生成内容")
    project.proposal_sections_json = _json_dumps(sections)
    project.status = "proposal_ready"
    _save_step(db, project, "proposal", sections)
    db.commit()
    db.refresh(project)
    return _project_to_response(project)


@router.post("/projects/{project_id}/exports/word")
async def export_grant_proposal_word(
    project_id: int,
    item: GrantStepAction,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin),
):
    project = await _get_project(db, project_id, item.token, current_user)
    buffer = _build_proposal_docx(project)
    filename = f"grant-proposal-{project_id}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
