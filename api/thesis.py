from __future__ import annotations

import io
import json
import random
import re
import time
import urllib.parse
from typing import List, Optional

import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException, Body, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from schemas.thesis import (
    ThesisCreate,
    ThesisOutlineRequest,
    ThesisOutlineSaveRequest,
    ThesisFullTextSaveRequest,
    ThesisDraftSaveRequest,
    ThesisFullTextRequest,
    ThesisRefineRequest,
    ThesisProjectResponse,
)
from utils.auth import get_optional_admin, check_permission, verify_token
from utils.docx_helpers import (
    add_latex_formatted_paragraph,
    add_latex_heading,
    add_hyperlink,
    _set_east_asian_font,
)
from database import get_db
from models import AdminUser, ThesisProject, ThesisStep
from api.dependencies import verify_service_access, deduct_token_quota
from services import ai_service, literature_service, ReferenceVerifier

router = APIRouter(prefix="/api/ai/thesis", tags=["thesis"])

# Cache for exported documents (project_id -> (buffer, timestamp))
_export_cache = {}
_export_cache_ttl = 300  # 5 minutes cache


@router.post("/create")
async def create_thesis_project(item: ThesisCreate, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, item.token, "ai")
    else:
        check_permission(current_user, "ai")

    project = ThesisProject(
        token_id=token_record.id if token_record else None,
        admin_id=current_user.id if current_user else None,
        title=f"关于 {item.topic} 的 {item.thesis_type or '论著'}",
        topic=item.topic,
        discipline=item.discipline,
        language=item.language,
        length=item.length,
        style=item.style,
        thesis_type=item.thesis_type
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return project


@router.post("/references/upload")
async def upload_and_verify_references(
    files: List[UploadFile] = File(...),
    skip_verification_indices: Optional[str] = Form(None),
    topic: Optional[str] = Form(None),
    discipline: Optional[str] = Form(None),
    keywords: Optional[str] = Form(None),
    style_example_idx: Optional[int] = Form(None),
    project_id: Optional[str] = Form(None),
    token: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """上传文献文件（PDF/Word），验证真实性并检查与研究主题的相关性"""
    import logging
    logging.warning(f"[UPLOAD] files={len(files)} token={token[:8] if token else None} project_id={project_id}")
    if not current_user:
        # Allow unauthenticated access with token if needed
        pass
    else:
        check_permission(current_user, "ai")

    if len(files) > 10:
        raise HTTPException(status_code=400, detail="最多只能上传 10 篇文献")

    # Parse skip_verification_indices
    skip_indices = set()
    if skip_verification_indices:
        try:
            skip_indices = {int(i.strip()) for i in skip_verification_indices.split(",") if i.strip()}
        except ValueError:
            pass

    # Build relevance context
    relevance_context = {}
    if topic:
        relevance_context["topic"] = topic
    if discipline:
        relevance_context["discipline"] = discipline
    if keywords:
        relevance_context["keywords"] = [k.strip() for k in keywords.split(",") if k.strip()]

    verifier = ReferenceVerifier()
    verified = []
    failed = []

    for idx, file in enumerate(files):
        filename = file.filename or "unknown"
        ext = filename.lower().split('.')[-1]

        if ext not in ('pdf', 'docx', 'doc'):
            failed.append({"filename": filename, "reason": "不支持的文件格式，仅支持 PDF 和 Word (.docx/.doc)"})
            continue

        try:
            content = await file.read()

            # 提取文本
            if ext == 'pdf':
                text = verifier.extract_text_from_pdf(content)
            else:
                text = verifier.extract_text_from_docx(content)

            if not text or len(text) < 100:
                failed.append({"filename": filename, "reason": "文件内容过短或无法提取有效文本"})
                continue

            # 提取元数据：使用大模型提取标题、作者、摘要等
            text_preview = text[:3000]
            metadata = await ReferenceVerifier.extract_metadata_with_ai(text_preview)
            title = metadata.get("title", "")
            doi = metadata.get("doi", "")

            if not title or len(title) < 5:
                failed.append({"filename": filename, "reason": "无法从文件中提取有效的论文标题"})
                continue

            # 保存摘要预览
            abstract = metadata.get("abstract", "")
            abstract_preview = abstract if abstract else text[:500].replace('\n', ' ').strip()

            # Check if this file should skip verification
            should_skip = idx in skip_indices

            # Check relevance to research topic/discipline
            if not should_skip:
                relevance = await ReferenceVerifier.check_relevance_with_ai(
                    title, abstract_preview, relevance_context
                )
                if not relevance["relevant"]:
                    failed.append({
                        "filename": filename,
                        "reason": f"文献与研究主题不相关（{relevance.get('reason', '')}），已自动忽略",
                        "is_irrelevant": True,
                    })
                    continue

                # 验证真实性：依次尝试 CrossRef、Europe PMC、arXiv
                verified_info = None

                # 如果有 DOI，先用 CrossRef 精确查询
                if doi:
                    verified_info = await verifier.verify_paper_crossref(title, doi)

                # CrossRef 按标题搜索（对中文论文使用 query.bibliographic 效果更好）
                if not verified_info:
                    verified_info = await verifier.verify_paper_crossref_by_title(title)

                # Europe PMC 搜索
                if not verified_info:
                    verified_info = await verifier.verify_paper_europepmc(title)

                # arXiv 搜索
                if not verified_info:
                    verified_info = await verifier.verify_paper_arxiv(title)

                if not verified_info:
                    failed.append({"filename": filename, "reason": f"未能通过 CrossRef/Europe PMC/arXiv 验证论文真实性（标题: {title[:60]}...），已自动忽略"})
                    continue
            else:
                # Skip verification: just use AI-extracted metadata
                verified_info = {
                    "title": title,
                    "authors": metadata.get("authors", []),
                    "doi": doi,
                    "year": metadata.get("year", ""),
                    "source": metadata.get("journal", ""),
                    "verified": False,
                    "database": "skipped",
                    "skipped_verification": True,
                }

            # 合并大模型提取的元数据
            verified_info["title"] = title or verified_info.get("title", "")
            if metadata.get("authors"):
                verified_info["authors"] = metadata["authors"]
            if abstract:
                verified_info["abstract_preview"] = abstract
            verified_info["filename"] = filename
            verified_info["upload_index"] = idx
            verified_info["is_style_example"] = (style_example_idx == idx)
            verified_info["year"] = metadata.get("year", verified_info.get("year", ""))
            verified_info["journal"] = metadata.get("journal", verified_info.get("journal", ""))
            verified.append(verified_info)
        except Exception as e:
            failed.append({"filename": filename, "reason": f"处理文件时出错: {str(e)}"})

    return {"verified": verified, "failed": failed, "total_uploaded": len(files), "verified_count": len(verified)}


@router.post("/outline")
async def generate_outline(req: ThesisOutlineRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        # Admin can see their own projects OR all if SuperAdmin
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found")

    # --- 处理用户上传的参考文献 ---
    user_refs = req.references or []
    style_example = req.style_example

    # 保存参考文献信息到项目
    if user_refs:
        project.reference_files = json.dumps(user_refs, ensure_ascii=False)
    if style_example:
        project.style_example_file = json.dumps(style_example, ensure_ascii=False)
    db.commit()

    user_ref_context = ""
    user_ref_bibliography = ""
    if user_refs:
        user_ref_context = "【用户上传的指定参考文献（已验证真实性或用户强制使用，必须在论文中引用并列入参考文献列表）】\n"
        for i, ref in enumerate(user_refs, 1):
            authors_str = ", ".join(ref.get("authors", [])[:3])
            if len(ref.get("authors", [])) > 3:
                authors_str += " et al."
            ref_line = f"{i}. {authors_str} ({ref.get('year', 'n.d.')}). {ref.get('title', '')}. {ref.get('source', '')}"
            if ref.get("doi"):
                ref_line += f". https://doi.org/{ref['doi']}"
            user_ref_context += f"  - {ref_line}\n"
            user_ref_bibliography += f"{ref_line}\n"
        user_ref_context += f"\n⚠️ 强制要求：以上 {len(user_refs)} 篇文献是用户指定的参考文献，必须全部在正文中通过文中引用格式（如 (Author et al., Year) 或 [数字]）进行引用，且必须全部出现在文末的参考文献列表中，一篇都不能遗漏。引用时应与正文论述自然结合，不可堆砌。\n"

    # --- Fetch Real Literature ---
    real_citations = await literature_service.search_literature(project.topic, max_results=15)
    lit_context = ""
    if real_citations:
        lit_context = "以下是从真实学术数据库中检索到的相关文献（可直接引用）：\n"
        lit_context += "\n".join([f"- {cit.get('formatted', str(cit))}" for cit in real_citations])

    # --- 基于论著类型的随机化策略（确保学术规范性） ---
    thesis_type = project.thesis_type or "原创性研究"

    # 1. 结构策略映射：确保不同论文类型采用符合学术规范的结构
    valid_strategies = {
        "原创性研究": [
            "采用标准 IMRAD 结构（引言-方法-结果-讨论），根据学科规范进行适度变体",
            "采用方法优先结构：突出方法学创新，随后展开应用与验证",
            "采用发现驱动结构：先呈现核心发现，再回溯方法与背景论证",
            "采用比较研究结构：方案 A 与方案 B 并行分析，最后综合比较",
            "采用多阶段研究结构：按实验或研究阶段分步报告与讨论"
        ],
        "综述": [
            "采用时间演进结构：历史回顾 → 发展阶段 → 当前进展 → 未来展望",
            "采用主题并列结构：围绕多个核心主题展开，最后综合比较",
            "采用方法分类结构：按研究方法或理论流派分类综述与评价",
            "采用争议聚焦结构：梳理主流观点 → 聚焦争议问题 → 分析不同立场 → 综合展望"
        ],
        "个案报告": [
            "采用标准个案报告结构：引言 → 病例描述 → 讨论 → 结论",
            "采用诊疗过程驱动结构：按接诊-诊断-治疗-随访时间线组织",
            "采用鉴别诊断结构：围绕主要鉴别诊断展开分析与排除"
        ],
        "Meta 分析": [
            "采用标准 PRISMA 规范结构：引言 → 方法 → 结果 → 讨论",
            "采用问题分层结构：按不同研究问题或亚组分析逐层报告"
        ],
        "致编辑信": [
            "采用论点-证据-建议结构：简明提出观点，辅以文献证据，最后给出建议"
        ]
    }

    # 2. 行文风格映射：排除不符合特定学科规范的风格
    valid_styles = {
        "原创性研究": [
            "严谨客观的实证风格，以数据和方法为核心驱动",
            "问题导向风格，围绕核心研究问题层层深入",
            "分析论证风格，注重实验结果与理论预期的对照"
        ],
        "综述": [
            "批判性分析风格，注重理论对话和学术争鸣",
            "比较综合风格，强调不同观点和方法的对比融合",
            "历史演进风格，展现理论或技术的发展脉络",
            "理论建构风格，注重概念框架的梳理与整合"
        ],
        "个案报告": [
            "临床叙事风格，客观详实记录诊疗全过程",
            "教学反思风格，侧重病例的临床启示和学习要点"
        ],
        "Meta 分析": [
            "严谨定量风格，以效应量和统计检验为核心",
            "方法学评价风格，注重纳入研究的质量评估与偏倚分析"
        ],
        "致编辑信": [
            "简明专业风格，语言礼貌、精炼、切中要害"
        ]
    }

    selected_strategy = random.choice(valid_strategies.get(thesis_type, valid_strategies["原创性研究"]))
    selected_style = random.choice(valid_styles.get(thesis_type, valid_styles["原创性研究"]))

    prompt = f"你是一位顶尖的学术导师。请为以下主题生成一份详细的论著提纲：\n"
    prompt += f"类型：{thesis_type}\n"
    prompt += f"主题：{project.topic}\n学科：{project.discipline}\n篇幅要求：{project.length}\n语言：{'中文' if project.language == 'zh' else '英文'}\n"
    prompt += f"【本次生成推荐采用】结构策略：{selected_strategy}\n"
    prompt += f"【本次生成推荐采用】行文风格：{selected_style}\n"
    if req.requirements:
        prompt += f"用户特别要求：{req.requirements}\n"

    if lit_context:
        prompt += f"\n{lit_context}\n"
    if user_ref_context:
        prompt += f"\n{user_ref_context}\n"
    if user_ref_bibliography:
        prompt += f"\n{user_ref_bibliography}\n"

    prompt += """
【极其重要的格式与内容要求】
1. 【绝对禁止】使用"第一章"、"第二章"、"第 1 章"、"Chapter 1"等任何非 Markdown 标题格式。
2. 【绝对禁止】在章节标题前后或内部添加任何额外标签，如"**核心章节**"、"**重点**"、"[关键]"、"*重要*"等。章节标题必须纯净，只包含标题文字。
3. 【必须使用】Markdown 标题符号组织全文层级结构：
   - 一级标题使用 `# 标题名称` 格式
   - 二级标题使用 `## 标题名称` 格式
   - 三级标题使用 `### 标题名称` 格式
   - 严禁使用纯数字编号（如 "1.", "1.1"）作为标题
4. 提纲的每个章节必须明确指向基于提供的真实文献可以进行深入分析和综合的具体学术问题。
5. 严禁生成空洞的框架式提纲。必须在每个二级/三级标题中明确指出该部分将分析的具体内容、将对比的具体文献、将探讨的具体研究问题。
6. 如果是 Meta 分析或综述，必须在提纲中明确指出将提取和整合哪些具体文献的哪些具体数据/结果。
7. 提纲必须体现深度批判性分析和学术综合能力，而非简单的结构罗列。
8. 章节标题必须纯净，严禁包含任何额外的 Markdown 标记（如 ** 或 _）、方括号标签（如 [核心]）、或任何形式的加重/标记符号。
9. 章节正文内容必须是纯文本，严禁在正文内部使用加粗（**）、斜体（_）或其他 Markdown 格式标记。
10. 直接输出提纲内容，严禁任何开场白、说明文字或结束语。
11. 不要输出"以下是..."、"好的..."、"我将..."等对话式文本。
12. 请根据具体的学科领域、研究类型和上方推荐的结构策略，生成详细、专业且结构独特的章节内容。
13. 公式必须使用 LaTeX 格式，块级公式使用 `$$ ... $$`，行内公式使用 `$ ... $`。

【结构多样性指导】
不要机械套用固定模板。根据研究主题的特点，你可以灵活选择和组合以下结构元素：

**原创性研究的常见结构变体**：
- 标准 IMRAD：引言 → 方法 → 结果 → 讨论
- 方法优先：方法（重点）→ 应用案例 → 结果验证 → 讨论
- 发现驱动：关键发现 → 方法回顾 → 结果展开 → 综合讨论
- 比较研究：研究设计 → 方案 A 分析 → 方案 B 分析 → 比较讨论
- 多阶段研究：阶段一 → 阶段二 → 阶段三 → 综合讨论

**综述论文的常见结构变体**：
- 时间演进：历史回顾 → 发展阶段 → 当前进展 → 未来展望
- 主题并列：主题 A → 主题 B → 主题 C → 综合比较
- 方法分类：方法类 A 综述 → 方法类 B 综述 → 比较评价
- 争议聚焦：主流观点 → 争议问题 → 不同立场 → 综合评价

**结构组织的技巧**：
- 可以根据研究内容的逻辑关系，调整章节的先后顺序
- 可以在主要章节之间添加过渡性小节
- 可以根据篇幅要求，适当增减章节层级（2-4 级标题）
- 结果部分的小节命名应反映具体发现，而非简单的"结果 1"、"结果 2"

【学科特定结构要求】
根据学科领域的不同，请在提纲中体现该学科特有的章节元素：
- **临床医学**：伦理审批与知情同意、诊断标准（WHO/ICD 等）、临床试验注册（如适用）、不良事件/副作用报告、CONSORT 声明遵循
- **生物信息学**：数据库版本与工具、分析流程与参数设置、质控标准、参考基因组版本、代码可重复性（GitHub）、可视化方案
- **计算机科学**：系统/模型架构、实验环境与硬件配置、评估指标与基线比较、消融实验、统计显著性检验、数据集来源与划分
- **公共卫生**：研究设计类型（横断面/队列/病例对照）、抽样方法、问卷/量表信效度、响应率与无应答偏倚、STROBE 声明、政策建议
- **基础医学**：动物模型/细胞系描述、实验分组与对照、分子/细胞生物学技术、统计方法与样本量计算、ARRIVE 指南遵循
- **药学**：药物来源与纯度、给药方案与途径、动物品系与培养条件、药效学/药代动力学参数、毒理学评估、标准单位报告
"""

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    outline = ai_response["content"]

    # Deduct token quota if using service token
    if token_record and ai_response.get("total_tokens", 0) > 0:
        deduct_token_quota(db, token_record.id, ai_response["total_tokens"])

    # Save step
    step = ThesisStep(project_id=project.id, step_num=1, content=outline)
    db.add(step)
    project.current_step = 1
    db.commit()

    return {"outline": outline, "project_id": project.id, "citations": real_citations}


@router.post("/outline/save")
async def save_outline(req: ThesisOutlineSaveRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Save outline as a new step
    step = ThesisStep(project_id=project.id, step_num=1, content=req.outline)
    db.add(step)
    if project.current_step < 1:
        project.current_step = 1
    db.commit()

    return {"message": "提纲已保存", "project_id": project.id}


@router.post("/fulltext/save")
async def save_fulltext(req: ThesisFullTextSaveRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Save fulltext as a new step (step_num = 2)
    step = ThesisStep(project_id=project.id, step_num=2, content=req.content)
    db.add(step)
    if project.current_step < 2:
        project.current_step = 2
    db.commit()

    return {"message": "全文内容已保存", "project_id": project.id}


@router.post("/draft/save")
async def save_draft(req: ThesisDraftSaveRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Save outline as a step (step_num = 1)
    outline_step = ThesisStep(project_id=project.id, step_num=1, content=req.outline)
    db.add(outline_step)

    # If chapters are provided, also save them as part of the outline metadata
    if req.chapters:
        outline_step.metadata_info = json.dumps({"chapters": req.chapters})

    if project.current_step < 1:
        project.current_step = 1
    db.commit()

    return {"message": "初校草稿已保存", "project_id": project.id}


@router.post("/fulltext")
async def generate_fulltext(req: ThesisFullTextRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found")

    # --- 处理用户上传的参考文献 ---
    user_refs = req.references or []
    style_example = req.style_example

    # 如果请求中没有传参考文献，尝试从项目中加载之前保存的
    if not user_refs and project.reference_files:
        try:
            user_refs = json.loads(project.reference_files)
        except:
            user_refs = []
    if not style_example and project.style_example_file:
        try:
            style_example = json.loads(project.style_example_file)
        except:
            style_example = None

    user_ref_context = ""
    user_ref_bibliography = ""
    if user_refs:
        user_ref_context = "【用户指定必须引用的参考文献（已通过真实性验证或用户强制使用）——这些文献必须在正文中引用，并列入文末参考文献列表】\n"
        user_ref_bibliography = "\n【必须在文末参考文献列表中完整包含的文献列表（请按以下格式逐条复制，不得遗漏任何一条）】\n"
        for i, ref in enumerate(user_refs, 1):
            authors_str = ", ".join(ref.get("authors", [])[:3])
            if len(ref.get("authors", [])) > 3:
                authors_str += " et al."
            ref_line = f"{i}. {authors_str} ({ref.get('year', 'n.d.')}). {ref.get('title', '')}. {ref.get('source', '')}"
            if ref.get("doi"):
                ref_line += f". https://doi.org/{ref['doi']}"
            user_ref_context += f"  - {ref_line}\n"
            user_ref_bibliography += f"{ref_line}\n"
        user_ref_context += f"\n⚠️ 核心要求：以上 {len(user_refs)} 篇文献必须在正文中通过文中引用格式（如 (Author et al., Year) 或 [数字]）进行引用，且必须全部出现在文末的参考文献列表中，一篇都不能遗漏。引用时应与正文论述自然结合，不可堆砌。\n"
        user_ref_bibliography += f"\n⚠️ 注意：以上 {len(user_refs)} 篇文献是用户指定的参考文献，必须在文末参考文献列表中逐条列出，不得修改、不得遗漏。请将它们放在参考文献列表的最前面（编号 1-{len(user_refs)}）。"

    style_example_context = ""
    if style_example:
        abstract_preview = style_example.get("abstract_preview", "")[:300]
        style_example_context = f"【用户提供的学术范例文献】\n标题：{style_example.get('title', '')}\n"
        if abstract_preview:
            style_example_context += f"内容摘要：{abstract_preview}...\n"
        style_example_context += "请适当参照该范例的学术写作风格、段落组织方式和论证逻辑。注意：这不是要复制内容，而是学习其行文风格。\n"

    # --- Fetch Real Literature ---
    real_citations = await literature_service.search_literature(project.topic, max_results=15)
    lit_context = ""
    if real_citations:
        lit_context = "以下是从 PubMed 真实学术数据库中检索到的可引用文献（共15篇），请在正文中通过 (Author et al., Year) 的方式进行引用，并确保最后的参考文献列表包含这些文献：\n"
        lit_context += "\n".join([f"- {cit.get('formatted', str(cit))}" for cit in real_citations])

    prompt = f"根据以下提纲，生成一篇完整的学术论著：\n提纲：\n{req.outline}\n\n要求：\n文风：{req.style or project.style}\n语言：{'中文' if project.language == 'zh' else '英文'}\n篇幅：{project.length}\n"
    if lit_context:
        prompt += f"\n{lit_context}\n"
    if user_ref_context:
        prompt += f"\n{user_ref_context}\n"
    if style_example_context:
        prompt += f"\n{style_example_context}\n"

    # 将参考文献模板单独放在 prompt 末尾（靠近输出指令的位置，确保 LLM 不会忘记）
    if user_ref_bibliography:
        prompt += f"\n{user_ref_bibliography}\n"

    # --- 基于论著类型的随机化行文策略（确保学术规范性） ---
    thesis_type = project.thesis_type or "原创性研究"

    # 1. 段落展开策略映射
    valid_paragraph_strategies = {
        "原创性研究": [
            "论述展开：采用'论点-数据-分析-小结'的经典四段式，客观呈现实证结果",
            "论述展开：先提出研究发现，再回溯实验设计和统计方法进行论证",
            "论述展开：从具体实验数据出发，逐步推导并上升到理论层面的讨论"
        ],
        "综述": [
            "论述展开：采用'总-分-总'结构，首尾呼应，中间综合比较不同文献观点",
            "论述展开：以文献对话展开，指出各研究的共识与分歧，进行批判性评价",
            "论述展开：按理论发展脉络组织，梳理不同阶段的代表性工作和突破"
        ],
        "个案报告": [
            "论述展开：按时间顺序客观记录诊疗过程，辅以文献对比分析",
            "论述展开：以临床问题为导向，逐步展示诊断思路和鉴别过程"
        ],
        "Meta 分析": [
            "论述展开：严格遵循 PRISMA 规范，逐项报告检索、筛选、提取和综合过程",
            "论述展开：先呈现整体合并效应量，再分层展开亚组分析和敏感性分析"
        ],
        "致编辑信": [
            "论述展开：简明提出核心观点，辅以关键文献证据，最后给出专业建议"
        ]
    }

    # 2. 段落衔接策略（通用学术规范，适度随机）
    valid_transitions = [
        "段落衔接：使用逻辑连接词（然而、因此、此外、值得注意的是）自然过渡",
        "段落衔接：使用承上启下的过渡句，总结前文要点并引出下文论述",
        "段落衔接：通过概念递进和层次深化实现段落之间的自然流转",
        "段落衔接：使用对比和转折关系连接不同段落，突出论述的学术张力"
    ]

    # 3. 论据呈现策略映射
    valid_evidence_styles = {
        "原创性研究": [
            "论据呈现：优先使用实验数据、统计检验结果和图表支撑论点",
            "论据呈现：结合本研究结果与现有文献比较，验证发现的一致性"
        ],
        "综述": [
            "论据呈现：通过多文献综合比较的方式构建论据链条，展现学术全景",
            "论据呈现：以里程碑式核心文献为骨架，辅以近年最新研究进行论证",
            "论据呈现：使用表格或分类梳理不同研究的发现、方法和局限"
        ],
        "个案报告": [
            "论据呈现：以患者客观临床数据（检查、检验、影像）为核心证据",
            "论据呈现：结合文献中类似病例的诊疗结果进行对比论证"
        ],
        "Meta 分析": [
            "论据呈现：严格报告效应量、置信区间、异质性检验等定量统计指标",
            "论据呈现：通过森林图、漏斗图等可视化工具辅助论证"
        ],
        "致编辑信": [
            "论据呈现：引用最近发表的相关文献，简明扼要支持论点"
        ]
    }

    selected_paragraph_strategy = random.choice(valid_paragraph_strategies.get(thesis_type, valid_paragraph_strategies["原创性研究"]))
    selected_transition = random.choice(valid_transitions)
    selected_evidence = random.choice(valid_evidence_styles.get(thesis_type, valid_evidence_styles["原创性研究"]))

    prompt += f"\n【本次生成推荐采用】{selected_paragraph_strategy}\n"
    prompt += f"【本次生成推荐采用】{selected_transition}\n"
    prompt += f"【本次生成推荐采用】{selected_evidence}\n"

    # 根据论著类型添加特定的格式要求
    thesis_type = project.thesis_type or "原创性研究"
    type_specific_format = ""

    if thesis_type == "原创性研究":
        type_specific_format = """
【原创性研究格式要求】
- 必须包含：标题、摘要（背景、目的、方法、结果、结论五要素）、关键词（3-8个）、引言、材料与方法、结果、讨论、结论、参考文献
- 摘要采用结构化格式（各要素分段写）
- 关键词应反映论文核心内容，使用 MeSH 术语（如适用）
- 方法部分应足够详细以确保可重复性
- 结果部分应客观呈现，避免主观解释
- 讨论部分应与已有研究比较并分析局限性
"""
    elif thesis_type == "综述":
        type_specific_format = """
【综述论文格式要求】
- 必须包含：标题、摘要（非结构化格式）、关键词（3-8个）、引言、正文（按主题组织）、讨论与展望、结论、参考文献
- 摘要为段落式，概括综述目的、主要发现和结论
- 正文应按逻辑主题或时间顺序组织，而非简单罗列文献
- 必须对不同研究的发现进行批判性比较和综合
- 应明确指出当前研究的共识、争议点和知识空白
- 讨论部分应提出未来研究方向和理论/实践意义
"""
    elif thesis_type == "个案报告":
        type_specific_format = """
【个案报告格式要求】
- 必须包含：标题（应包含"个案报告"或"Case Report"字样）、摘要、关键词（3-5个）、引言、病例描述、讨论、结论、参考文献
- 病例描述应包含：患者人口学特征（匿名化）、主诉、现病史、既往史、体格检查、辅助检查、诊断、治疗和随访
- 讨论应将该病例与文献中类似病例比较，指出独特性或教学价值
- 应说明此个案报告的临床启示和学习要点
- 如适用，应包含伦理审批或患者知情同意说明
"""
    elif thesis_type == "Meta 分析":
        type_specific_format = """
【Meta 分析格式要求】
- 必须包含：标题（应包含"Meta 分析"字样）、摘要（结构化）、关键词、引言、方法（检索策略、纳入排除标准、数据提取、质量评价、统计分析）、结果（文献筛选、纳入研究特征、定量综合结果）、讨论、结论、参考文献
- 如遵循 PRISMA 指南，应在方法部分声明
- 必须详细描述文献检索策略（数据库、时间范围、检索词）
- 必须报告效应量（如 OR, RR, SMD 等）及其 95% 置信区间
- 必须报告异质性检验结果（I² 统计量、Cochran's Q 检验 p 值）
- 如进行亚组分析或敏感性分析，应报告结果
- 必须评估发表偏倚（如漏斗图、Egger's test）
- 所有数据必须来自提供的真实文献，绝不可虚构任何统计值
"""
    else:  # 致编辑信
        type_specific_format = """
【致编辑信格式要求】
- 必须包含：标题、正文（简明扼要）、参考文献
- 篇幅通常较短（500-1000字）
- 语言应礼貌、专业、简洁
- 应清晰表达观点或回应
- 如有引用，应适当引用支持论点
"""

    # 根据学科添加特定要求
    discipline = project.discipline or ""
    discipline_specific = ""

    if "临床" in discipline:
        discipline_specific = """
【临床医学学科特定要求】
- 如涉及患者数据，应说明伦理审批和知情同意
- 诊断标准应参考国际公认标准（如 WHO、ICD 等）
- 如为临床试验，应遵循 CONSORT 声明
- 统计方法应说明样本量计算依据
- 应报告不良反应或副作用（如适用）
- 参考文献格式应采用 Vancouver 格式（数字编号系统）：
  * 文内引用使用上标数字或方括号数字，如 [1] 或 ¹
  * 参考文献列表按引用顺序编号排列
  * 格式：作者. 题名. 期刊名. 年;卷(期):起止页码. DOI
  * 示例：Smith J, Johnson K, Williams R. Title of article. J Clin Med. 2023;15(2):123-135. https://doi.org/10.1234/example
"""
    elif "生物信息" in discipline or "生物" in discipline:
        discipline_specific = """
【生物信息学学科特定要求】
- 应说明使用的数据库版本和参数设置
- 应说明数据处理流程和质控标准
- 如涉及基因组数据，应说明参考基因组版本
- 应提供代码或分析流程的可重复性信息（如 GitHub 链接，可占位）
- 可视化图表应清晰标注坐标轴和统计信息
- 参考文献格式应采用 APA 格式：
  * 文内引用使用 (作者, 年份) 格式
  * 参考文献列表按作者姓氏字母顺序排列
  * 格式：作者. (年份). 题名. 期刊名, 卷(期), 起止页码. DOI链接
  * 示例：Smith, J., Johnson, K., & Williams, R. (2023). Title of article. Bioinformatics, 39(2), 123-135. https://doi.org/10.1234/example
"""
    elif "计算机" in discipline or "软件" in discipline or "AI" in discipline:
        discipline_specific = """
【计算机科学学科特定要求】
- 应详细描述算法、模型架构和超参数设置
- 实验设置应包含硬件环境、软件版本和依赖
- 应使用标准评估指标并与基线方法比较
- 应进行统计显著性检验（如 t 检验、ANOVA）
- 如使用公开数据集，应说明数据来源和划分方式
- 参考文献格式应采用 IEEE 格式：
  * 文内引用使用方括号数字，如 [1]
  * 参考文献列表按引用顺序编号排列
  * 格式：作者, "题名," 期刊名, 卷号, 期号, 起止页码, 年份. DOI
  * 示例：J. Smith, K. Johnson, and R. Williams, "Title of article," IEEE Trans. Comput., vol. 72, no. 3, pp. 123-135, 2023. https://doi.org/10.1234/example
"""
    elif "公共" in discipline or "卫生" in discipline or "流行病学" in discipline:
        discipline_specific = """
【公共卫生学科特定要求】
- 研究设计应明确说明是横断面、队列、病例对照还是实验性研究
- 应说明研究人群的纳入排除标准和抽样方法
- 如为调查研究，应说明问卷的信效度检验
- 应报告响应率和无应答偏倚评估
- 如遵循 STROBE 声明，应声明
- 参考文献格式应采用 APA 格式：
  * 文内引用使用 (作者, 年份) 格式
  * 参考文献列表按作者姓氏字母顺序排列
  * 格式：作者. (年份). 题名. 期刊名, 卷(期), 起止页码. DOI链接
"""
    elif "基础" in discipline and "医学" in discipline:
        discipline_specific = """
【基础医学学科特定要求】
- 应详细描述动物模型或细胞系（品系、来源、培养条件）
- 应说明实验分组和对照组设计
- 分子/细胞生物学技术应说明实验步骤和试剂
- 应说明样本量计算和统计检验方法
- 如为动物实验，应遵循 ARRIVE 指南并说明伦理审批
- 参考文献格式应采用 Vancouver 格式：
  * 文内引用使用上标数字或方括号数字
  * 参考文献列表按引用顺序编号排列
  * 格式：作者. 题名. 期刊名. 年;卷(期):起止页码. DOI
"""
    elif "药学" in discipline or "药物" in discipline:
        discipline_specific = """
【药学学科特定要求】
- 应说明药物来源、纯度和给药方式
- 实验动物或细胞应说明品系、来源和培养条件
- 应遵循 ARRIVE 指南（动物实验）
- 药效学/药代动力学参数应使用标准单位
- 应说明统计分析方法和显著性水平
- 参考文献格式应采用 Vancouver 格式：
  * 文内引用使用上标数字或方括号数字
  * 参考文献列表按引用顺序编号排列
  * 格式：作者. 题名. 期刊名. 年;卷(期):起止页码. DOI
"""
    else:
        discipline_specific = """
【通用学科要求】
- 参考文献格式应采用 APA 格式（第7版）：
  * 文内引用使用 (作者, 年份) 格式，多位作者首次使用列出全部，后续使用 et al. 或 等
  * 参考文献列表按作者姓氏字母顺序排列
  * 格式：作者. (年份). 题名. 期刊名, 卷(期), 起止页码. DOI链接（如有）
  * 示例：Smith, J., Johnson, K., & Williams, R. (2023). Title of article. Journal Name, 15(2), 123-135. https://doi.org/10.1234/example
  * 期刊名使用标准缩写或全称
  * DOI 必须作为可点击链接提供：https://doi.org/DOI号码
"""

    prompt += f"""
【极其重要的核心要求】
1. 严禁生成框架式或占位符式的文章！绝对不要使用"此处应包含具体数据"、"如表X所示"、"详见图Y"等提示性文字。
2. 严禁杜撰、编造任何研究数据、实验结果、统计值（如 p 值、OR 值、样本量、效应量等）或临床发现。
3. 必须基于真实的学术文献（已在上方提供）进行深度综合、逻辑推理和学术论述，生成内容充实、数据详实、论证严谨的真正学术论文。
4. 如果是 Meta 分析或综述类文章，必须引用并整合上述真实文献中的具体研究结果、样本量、效应量等真实数据。如果没有足够的定量数据，则应进行深度的定性综合和批判性分析，绝不可虚构任何数值或统计结果。
5. 所有引用的文献必须且只能来自上方提供的真实文献列表。

【格式与输出要求】
1. 直接输出论文正文内容，不要有任何开场白、说明文字或结束语
2. 不要输出"好的..."、"遵照您的要求..."、"我将..."等对话式文本
3. 从标题或摘要开始直接输出，不要有额外的解释
4. 请在文中适当位置使用上述学科指定的引用格式引用上述真实文献
5. 在文末必须提供完整的参考文献列表（References），包含所有引用的文献。如果用户提供了指定参考文献，这些文献必须全部出现在参考文献列表中，一篇都不能遗漏
6. 参考文献格式必须严格遵循上述学科指定的格式要求，确保：
   - 所有作者姓名完整且格式正确
   - 年份、标题、期刊名准确无误
   - DOI 必须以完整的可点击链接形式提供：https://doi.org/DOI号码
   - 切勿编造任何文献信息或 DOI
7. 绝对不要编造不存在的文献或 DOI
8. 所有引用必须来自上述提供的真实文献列表
9. 【极其重要】必须使用 Markdown 标题语法组织全文结构：
   - 一级标题使用 `# 1. 标题名称` 格式
   - 二级标题使用 `## 1.1 标题名称` 格式
   - 三级标题使用 `### 1.1.1 标题名称` 格式
   - 所有章节编号必须按层级递进：1 → 1.1 → 1.1.1 → 1.2 → 2 → 2.1 等
   - 每个章节必须有对应的层级标题
   - 示例结构：
     ```
     # 1. 引言

     引言内容...

     ## 1.1 研究背景

     背景内容...

     ### 1.1.1 国内研究现状

     国内研究现状内容...

     # 2. 方法

     方法内容...

     ## 2.1 实验设计

     实验设计内容...
     ```
10. {type_specific_format.strip()}
11. {discipline_specific.strip() if discipline_specific else '遵循所属学科的一般学术规范和格式要求。'}
12. 使用专业、严谨的学术语言
13. 确保逻辑连贯，各部分之间有良好的衔接
14. 适当使用学术术语，但避免过度使用
15. 数据和结果描述应准确、客观
16. 公式必须使用 LaTeX 格式，块级公式使用 `$$ ... $$`，行内公式使用 `$ ... $`。

【行文多样性要求】
- 避免千篇一律的句式，灵活运用长短句结合、主动被动语态交替、不同句型变化
- 论述方式可以多样化：可直接阐述、可通过对比引入、可先提出问题再解答、可先陈述共识再指出分歧
- 避免每个段落都以相同方式开头（如总是"研究表明..."或"近年来..."）
- 使用丰富的学术表达词汇：不仅使用"表明"，也可使用"揭示"、"提示"、"印证"、"佐证"、"阐明"等
- 在讨论部分可以采用不同角度的分析：与理论预期比较、与同类研究比较、从方法学角度分析、从临床应用角度讨论
"""

    # 最后附加参考文献强制要求（放在 prompt 最末尾，确保 LLM 不会忽略）
    if user_refs:
        prompt += f"\n\n🔴 最终提醒：在输出论文的最后，你必须提供一个完整的参考文献列表（References）。以下 {len(user_refs)} 篇文献必须一字不差地逐条列在参考文献列表中（可放在其他文献之前），一篇都不能少：\n"
        for i, ref in enumerate(user_refs, 1):
            authors_str = ", ".join(ref.get("authors", [])[:3])
            if len(ref.get("authors", [])) > 3:
                authors_str += " et al."
            ref_line = f"{i}. {authors_str} ({ref.get('year', 'n.d.')}). {ref.get('title', '')}. {ref.get('source', '')}"
            if ref.get("doi"):
                ref_line += f". https://doi.org/{ref['doi']}"
            prompt += f"{ref_line}\n"

    # 随机化 temperature 以增加生成多样性 (0.7-0.9)
    generation_temperature = round(random.uniform(0.7, 0.9), 1)

    ai_response = await ai_service.chat_completion([{"role": "system", "content": "你是一位专业的学术写作者，擅长根据研究主题灵活选择最合适的论述风格和文章架构。直接输出论文内容，不要有任何对话式文本。"}, {"role": "user", "content": prompt}], temperature=generation_temperature)
    fulltext = ai_response["content"]

    # --- 后处理：确保用户指定的文献一定出现在参考文献列表中 ---
    if user_refs:
        # 生成用户文献的参考文献条目
        user_bib_lines = []

        def _find_ref_section(text):
            """查找参考文献章节，返回 (start, end) 或 None"""
            # 策略 1: 匹配 # 开头的 Markdown 标题（最常见）
            patterns_markdown = [
                r'^#{1,3}\s*\d*[\.\s]*(?:参考文献|References|REF)\s*$',  # # 参考文献
                r'^#{1,3}\s*(?:参考文献|References|REF)\s*$',             # # 参考文献（无编号）
            ]
            for pat in patterns_markdown:
                m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
                if m:
                    section_start = m.end()
                    next_heading = re.search(r'\n#{1,3}\s', text[section_start:])
                    section_end = section_start + next_heading.start() if next_heading else len(text)
                    return (section_start, section_end)

            # 策略 2: 匹配纯文本标题（无 # 号，但在行首）
            plain_patterns = [
                r'^(?:参考文献|References|REF)\s*$',
                r'^\d+[\.\s]*(?:参考文献|References|REF)\s*$',
                r'^[（\(][一二三四五六七八九十\d]+[）\)]\s*(?:参考文献|References|REF)\s*$',
            ]
            for pat in plain_patterns:
                m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
                if m:
                    section_start = m.end()
                    # 找到下一个章节标题（带编号的行或 # 开头的行）
                    next_heading = re.search(r'\n(?:#{1,3}\s|\d+[\.\s、]|第[一二三四五六七八九十百\d]+[章节篇部])', text[section_start:])
                    section_end = section_start + next_heading.start() if next_heading else len(text)
                    return (section_start, section_end)

            # 策略 3: 匹配加粗的标题
            bold_match = re.search(r'\*{1,2}(?:参考文献|References|REF)\*{1,2}', text, re.IGNORECASE)
            if bold_match:
                section_start = bold_match.end()
                next_heading = re.search(r'\n(?:#{1,3}\s|\d+[\.\s、]|\*{1,2}\d)', text[section_start:])
                section_end = section_start + next_heading.start() if next_heading else len(text)
                return (section_start, section_end)

            return None

        ref_section = _find_ref_section(fulltext)

        if ref_section:
            ref_section_start, ref_section_end = ref_section
            ref_section_content = fulltext[ref_section_start:ref_section_end]

            # 找出已有参考文献的最大编号
            existing_refs = re.findall(r'^(\d+)\.\s', ref_section_content, re.MULTILINE)
            next_num = int(max(existing_refs)) + 1 if existing_refs else 1

            # 生成用户文献条目
            for i, ref in enumerate(user_refs):
                authors_str = ", ".join(ref.get("authors", [])[:3])
                if len(ref.get("authors", [])) > 3:
                    authors_str += " et al."
                ref_line = f"{next_num + i}. {authors_str} ({ref.get('year', 'n.d.')}). {ref.get('title', '')}. {ref.get('source', '')}"
                if ref.get("doi"):
                    ref_line += f". https://doi.org/{ref['doi']}"
                user_bib_lines.append(ref_line)

            user_bib_block = "\n".join(user_bib_lines)

            # 在参考文献章节末尾追加（在下一个标题之前）
            insert_text = "\n" + user_bib_block + "\n"
            fulltext = fulltext[:ref_section_end] + insert_text + fulltext[ref_section_end:]
        else:
            # 没有参考文献部分，在文末追加
            next_num = 1
            for i, ref in enumerate(user_refs):
                authors_str = ", ".join(ref.get("authors", [])[:3])
                if len(ref.get("authors", [])) > 3:
                    authors_str += " et al."
                ref_line = f"{next_num + i}. {authors_str} ({ref.get('year', 'n.d.')}). {ref.get('title', '')}. {ref.get('source', '')}"
                if ref.get("doi"):
                    ref_line += f". https://doi.org/{ref['doi']}"
                user_bib_lines.append(ref_line)
            fulltext += "\n\n# 参考文献\n" + "\n".join(user_bib_lines) + "\n"

    # Deduct token quota if using service token
    if token_record and ai_response.get("total_tokens", 0) > 0:
        deduct_token_quota(db, token_record.id, ai_response["total_tokens"])

    # Save step
    step = ThesisStep(project_id=project.id, step_num=2, content=fulltext)
    db.add(step)
    project.current_step = 2
    db.commit()

    return {"fulltext": fulltext, "project_id": project.id, "citations": real_citations}


@router.get("/projects", response_model=List[ThesisProjectResponse])
async def list_thesis_projects(token: Optional[str] = None, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        if not token:
            raise HTTPException(status_code=401, detail="Access token required")
        # Just verify token, don't deduct quota for listing projects
        token_record = verify_token(token, db, "ai")
        return db.query(ThesisProject).filter(ThesisProject.token_id == token_record.id).all()

    check_permission(current_user, "ai")
    # If SuperAdmin, show all. Otherwise show only projects created by this admin
    if current_user.group and current_user.group.name == "SuperAdmin":
        return db.query(ThesisProject).all()
    return db.query(ThesisProject).filter(ThesisProject.admin_id == current_user.id).all()


@router.get("/{project_id}/steps")
async def get_thesis_steps(project_id: int, token: Optional[str] = None, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found")
    return db.query(ThesisStep).filter(ThesisStep.project_id == project_id).order_by(ThesisStep.step_num.asc(), ThesisStep.created_at.desc()).all()


@router.get("/{project_id}/references")
async def get_project_references(project_id: int, token: Optional[str] = None, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    """Get all references for a project — both user-uploaded and auto-retrieved."""
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found")

    # 1. User-uploaded references
    uploaded_refs = []
    if project.reference_files:
        try:
            uploaded_refs = json.loads(project.reference_files)
        except (json.JSONDecodeError, TypeError):
            uploaded_refs = []

    # 2. Auto-retrieved literature
    real_citations = await literature_service.search_literature(project.topic, max_results=15)

    # 3. Style example
    style_example = None
    if project.style_example_file:
        try:
            style_example = json.loads(project.style_example_file)
        except (json.JSONDecodeError, TypeError):
            style_example = None

    return {
        "project_id": project_id,
        "uploaded": uploaded_refs,
        "retrieved": real_citations,
        "style_example": style_example,
        "uploaded_count": len(uploaded_refs),
        "retrieved_count": len(real_citations),
    }


@router.post("/refine")
async def refine_thesis_part(req: ThesisRefineRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, req.token, "ai")
    else:
        check_permission(current_user, "ai")

    # Check ownership
    query = db.query(ThesisProject).filter(ThesisProject.id == req.project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found or access denied")

    instructions = {
        "polish": "润色这段文字，使其更具学术专业性",
        "expand": "对这段文字进行扩写，增加细节和深度",
        "shorten": "对这段文字进行精简，保留核心意思",
        "reduce_similarity": "对这段学术论著进行深度降重，通过句式重构、同义替换、语态转换等手段有效降低查重率",
        "rewrite": "对这段文字进行全面改写，用全新表达方式呈现相同内容",
        "grammar": "纠正这段文字中的语法和拼写错误",
        "proofread": "对这段文字进行终极校对，检查语法、逻辑和引用规范",
        "style_change": "调整这段文字的文风，使其更符合指定的写作风格",
        "translate": "将这段文字进行中英互译"
    }

    instruction_text = instructions.get(req.instruction, req.instruction)
    prompt = f"针对以下学术论著片段，请执行以下操作：{instruction_text}。\n\n待处理文本：\n{req.text}"

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    result = ai_response["content"]

    # Deduct token quota if using service token
    if token_record and ai_response.get("total_tokens", 0) > 0:
        deduct_token_quota(db, token_record.id, ai_response["total_tokens"])

    # We don't necessarily save every small refinement as a new step unless requested,
    # but for now we just return the result.
    return {"result": result}


@router.post("/outline/export")
async def export_outline_word(req: dict = Body(...), db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    """导出提纲为 Word 文档"""
    # Authorization
    token = req.get("token")
    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, token, "ai")
    else:
        check_permission(current_user, "ai")

    title = req.get("title", "论文提纲")
    content_md = req.get("content", "")

    if not content_md:
        raise HTTPException(status_code=400, detail="No content to export")

    doc = docx.Document()

    # ===== 设置全局样式 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)  # 五号字
    if style.element.rPr is None:
        style.element._rPr = docx.oxml.OxmlElement('w:rPr')
    rFonts = style.element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        style.element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')  # 中文宋体

    # 段落格式
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 设置标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = 'Times New Roman'
        if heading_style.element.rPr is None:
            heading_style.element._rPr = docx.oxml.OxmlElement('w:rPr')
        rF = heading_style.element.rPr.find(qn('w:rFonts'))
        if rF is None:
            rF = docx.oxml.OxmlElement('w:rFonts')
            heading_style.element.rPr.append(rF)
        rF.set(qn('w:eastAsia'), '黑体')  # 中文标题用黑体
        if level == 1:
            heading_font.size = Pt(16)
            heading_font.bold = True
        elif level == 2:
            heading_font.size = Pt(14)
            heading_font.bold = True
        else:
            heading_font.size = Pt(12)
            heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    # ===== 添加标题 =====
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_font = title_para.runs[0].font
    title_font.name = 'Times New Roman'
    if title_para.runs[0].element.rPr is None:
        title_para.runs[0].element._r.get_or_add_rPr()
    rFonts = title_para.runs[0].element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = docx.oxml.OxmlElement('w:rFonts')
        title_para.runs[0].element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '黑体')
    title_font.size = Pt(22)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)

    # ===== 辅助函数：添加带有内联格式（加粗、斜体）的段落 =====
    def add_formatted_paragraph(doc, text, style='Normal'):
        """添加带有内联格式（加粗、斜体、行内代码）的段落"""
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        remaining = text
        while remaining:
            # 加粗 **text**
            if remaining.startswith('**'):
                end = remaining.find('**', 2)
                if end != -1:
                    run = p.add_run(remaining[2:end])
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    if run.element.rPr is None:
                        run.element._r.get_or_add_rPr()
                    rF = run.element.rPr.find(qn('w:rFonts'))
                    if rF is None:
                        rF = docx.oxml.OxmlElement('w:rFonts')
                        run.element.rPr.append(rF)
                    rF.set(qn('w:eastAsia'), '宋体')
                    remaining = remaining[end+2:]
                    continue
            # 斜体 *text*
            elif remaining.startswith('*') and not remaining.startswith('**'):
                end = remaining.find('*', 1)
                if end != -1:
                    run = p.add_run(remaining[1:end])
                    run.italic = True
                    run.font.name = 'Times New Roman'
                    if run.element.rPr is None:
                        run.element._r.get_or_add_rPr()
                    rF = run.element.rPr.find(qn('w:rFonts'))
                    if rF is None:
                        rF = docx.oxml.OxmlElement('w:rFonts')
                        run.element.rPr.append(rF)
                    rF.set(qn('w:eastAsia'), '宋体')
                    remaining = remaining[end+1:]
                    continue
            # 行内代码 `text`
            elif remaining.startswith('`'):
                end = remaining.find('`', 1)
                if end != -1:
                    run = p.add_run(remaining[1:end])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    if run.element.rPr is None:
                        run.element._r.get_or_add_rPr()
                    rF = run.element.rPr.find(qn('w:rFonts'))
                    if rF is None:
                        rF = docx.oxml.OxmlElement('w:rFonts')
                        run.element.rPr.append(rF)
                    rF.set(qn('w:eastAsia'), '宋体')
                    remaining = remaining[end+1:]
                    continue

            # 查找下一个格式标记
            next_markers = []
            for marker in ['**', '*', '`']:
                idx = remaining.find(marker)
                if idx != -1:
                    next_markers.append((idx, marker))

            if next_markers:
                next_markers.sort()
                idx, marker = next_markers[0]
                if idx > 0:
                    run = p.add_run(remaining[:idx])
                    run.font.name = 'Times New Roman'
                    if run.element.rPr is None:
                        run.element._r.get_or_add_rPr()
                    rF = run.element.rPr.find(qn('w:rFonts'))
                    if rF is None:
                        rF = docx.oxml.OxmlElement('w:rFonts')
                        run.element.rPr.append(rF)
                    rF.set(qn('w:eastAsia'), '宋体')
                remaining = remaining[idx:]
            else:
                run = p.add_run(remaining)
                run.font.name = 'Times New Roman'
                if run.element.rPr is None:
                    run.element._r.get_or_add_rPr()
                rF = run.element.rPr.find(qn('w:rFonts'))
                if rF is None:
                    rF = docx.oxml.OxmlElement('w:rFonts')
                    run.element.rPr.append(rF)
                rF.set(qn('w:eastAsia'), '宋体')
                remaining = ''
        return p

    # ===== 解析 Markdown 内容 =====
    lines_md = content_md.split('\n')
    for line in lines_md:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue

        p = None
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ '):
            add_formatted_paragraph(doc, stripped[2:], style='List Bullet')
            continue
        elif stripped and stripped[0].isdigit() and '. ' in stripped[:5]:
            text = stripped.split('. ', 1)[1] if '. ' in stripped else stripped
            add_formatted_paragraph(doc, text, style='List Number')
            continue
        else:
            add_formatted_paragraph(doc, stripped)
            continue

        # 为标题段落也设置字体
        if p:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.color.rgb = RGBColor(0, 0, 0)
                if run.element.rPr is None:
                    run.element._r.get_or_add_rPr()
                rF = run.element.rPr.find(qn('w:rFonts'))
                if rF is None:
                    rF = docx.oxml.OxmlElement('w:rFonts')
                    run.element.rPr.append(rF)
                rF.set(qn('w:eastAsia'), '宋体')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # URL-encode filename for Content-Disposition header (RFC 5987)
    encoded_filename = urllib.parse.quote(f"{title}_outline.docx")

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@router.get("/{project_id}/export")
async def export_thesis(project_id: int, token: Optional[str] = None, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    # Check cache first (if not forcing refresh)
    current_time = time.time()
    if project_id in _export_cache:
        cached_buffer, cached_time, cached_title = _export_cache[project_id]
        if current_time - cached_time < _export_cache_ttl:
            # Return cached version
            cached_buffer.seek(0)
            new_buffer = io.BytesIO(cached_buffer.read())
            new_buffer.seek(0)

            encoded_filename = urllib.parse.quote(f"{cached_title}.docx")
            return StreamingResponse(
                new_buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
                }
            )

    token_record = None
    if not current_user:
        token_record = await verify_service_access(db, token, "ai")
    else:
        check_permission(current_user, "ai")

    query = db.query(ThesisProject).filter(ThesisProject.id == project_id)
    if token_record:
        query = query.filter(ThesisProject.token_id == token_record.id)
    elif current_user:
        if current_user.group and current_user.group.name != "SuperAdmin":
            query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project: raise HTTPException(status_code=404, detail="Project not found")

    # Get the latest step content (full text)
    latest_step = db.query(ThesisStep).filter(ThesisStep.project_id == project_id, ThesisStep.step_num == 2).order_by(ThesisStep.created_at.desc()).first()
    if not latest_step:
        latest_step = db.query(ThesisStep).filter(ThesisStep.project_id == project_id).order_by(ThesisStep.step_num.desc(), ThesisStep.created_at.desc()).first()

    if not latest_step: raise HTTPException(status_code=400, detail="No content to export")

    # Helper function to safely set East Asian font on runs/styles
    def set_east_asian_font(font_obj, font_name):
        """Safely set the East Asian font for a docx Font object"""
        rpr = font_obj._element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = docx.oxml.OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:eastAsia'), font_name)

    # Create Word Document with academic formatting
    doc = docx.Document()

    # ===== 设置全局样式 =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)
    set_east_asian_font(font, '宋体')  # 中文宋体

    # 段落格式
    pf = style.paragraph_format
    pf.space_before = docx.shared.Pt(0)
    pf.space_after = docx.shared.Pt(6)
    pf.line_spacing = 1.5  # 1.5 倍行距
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐

    # 设置标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = 'Times New Roman'
        set_east_asian_font(heading_font, '黑体')  # 中文标题用黑体
        if level == 1:
            heading_font.size = docx.shared.Pt(16)
            heading_font.bold = True
        elif level == 2:
            heading_font.size = docx.shared.Pt(14)
            heading_font.bold = True
        else:
            heading_font.size = docx.shared.Pt(12)
            heading_font.bold = True
        heading_style.paragraph_format.space_before = docx.shared.Pt(12)
        heading_style.paragraph_format.space_after = docx.shared.Pt(6)

    # ===== 添加论文标题 =====
    title_para = doc.add_heading(project.title, level=0)
    title_font = title_para.runs[0].font
    title_font.name = 'Times New Roman'
    set_east_asian_font(title_para.runs[0], '黑体')
    title_font.size = docx.shared.Pt(22)
    title_font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===== 添加基本信息（可选，作为文档属性） =====
    doc.add_paragraph('')  # 空行
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"学科：{project.discipline}  |  类型：{project.thesis_type}  |  篇幅：{project.length}")
    info_run.font.size = docx.shared.Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    set_east_asian_font(info_run, '宋体')

    # ===== 添加分页符（可选，如果需要正文从新页开始） =====
    # doc.add_page_break()

    # ===== 解析 Markdown 内容并转换为 Word 格式 =====
    content_md = latest_step.content
    lines_md = content_md.split('\n')

    in_list = False
    in_code_block = False
    current_list_type = None  # 'bullet', 'numbered', or None

    for line in lines_md:
        stripped = line.strip()

        # 代码块
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            in_list = False
            current_list_type = None
            continue
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = docx.shared.Pt(0)
            p.paragraph_format.space_after = docx.shared.Pt(6)
            run = p.add_run(stripped)
            run.font.name = 'Courier New'
            run.font.size = docx.shared.Pt(10)
            p.paragraph_format.left_indent = docx.shared.Inches(0.5)
            continue

        # 标题 - 可能包含 LaTeX
        if stripped.startswith('### '):
            add_latex_heading(doc, stripped[4:], level=3)
            in_list = False
            current_list_type = None
        elif stripped.startswith('## '):
            add_latex_heading(doc, stripped[3:], level=2)
            in_list = False
            current_list_type = None
        elif stripped.startswith('# '):
            add_latex_heading(doc, stripped[2:], level=1)
            in_list = False
            current_list_type = None
        # 无序列表
        elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ '):
            in_list = True
            current_list_type = 'bullet'
            text = stripped[2:] if stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ ') else stripped
            add_latex_formatted_paragraph(doc, text, style='List Bullet')
        # 有序列表
        elif stripped and len(stripped) > 2 and stripped[0].isdigit() and '. ' in stripped[:5]:
            in_list = True
            current_list_type = 'numbered'
            parts = stripped.split('. ', 1)
            text = parts[1] if len(parts) > 1 else stripped
            add_latex_formatted_paragraph(doc, text, style='List Number')
        # 引用块
        elif stripped.startswith('>'):
            in_list = False
            current_list_type = None
            p = add_latex_formatted_paragraph(doc, stripped[1:].strip())
            p.italic = True
            p.paragraph_format.left_indent = docx.shared.Inches(0.5)
        # 空行
        elif stripped == '':
            in_list = False
            current_list_type = None
            doc.add_paragraph('')
        # 普通段落 - 使用 LaTeX 感知函数
        else:
            in_list = False
            current_list_type = None
            add_latex_formatted_paragraph(doc, stripped)

    # 保存并返回
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Cache the generated buffer
    _export_cache[project_id] = (io.BytesIO(buffer.getvalue()), current_time, project.title)

    # Clean up old cache entries
    expired_keys = [k for k, (buf, ts, title) in _export_cache.items() if current_time - ts >= _export_cache_ttl]
    for k in expired_keys:
        del _export_cache[k]

    # Properly encode filename for Content-Disposition header
    encoded_filename = urllib.parse.quote(f"{project.title}.docx")

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


@router.post("/{project_id}/validate-references")
async def validate_thesis_references(project_id: int, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    """
    Validate that references in the thesis text match the provided citation list.

    This endpoint analyzes a thesis project to ensure that:
    1. In-text citations are properly formatted
    2. The reference list is present
    3. Provided citations are actually used in the text
    """
    token = None  # For future auth
    query = db.query(ThesisProject).filter(ThesisProject.id == project_id)
    if current_user and current_user.group and current_user.group.name != "SuperAdmin":
        query = query.filter(ThesisProject.admin_id == current_user.id)

    project = query.first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Get the full text (step 2)
    fulltext_step = db.query(ThesisStep).filter(
        ThesisStep.project_id == project_id,
        ThesisStep.step_num == 2
    ).order_by(ThesisStep.created_at.desc()).first()

    if not fulltext_step:
        raise HTTPException(status_code=404, detail="No full text found for this project")

    fulltext = fulltext_step.content

    # Get citations from metadata or search fresh
    citations = []
    if fulltext_step.metadata_info:
        try:
            metadata = json.loads(fulltext_step.metadata_info)
            citations = metadata.get("citations", [])
        except:
            pass

    # If no citations in metadata, search fresh
    if not citations:
        citations = await literature_service.search_literature(project.topic, max_results=12)

    # Validate references
    validation_result = literature_service.validate_references(fulltext, citations)

    return {
        "project_id": project_id,
        "project_title": project.title,
        "validation": validation_result
    }
