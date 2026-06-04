from __future__ import annotations

import io
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from schemas.literature import (
    LiteratureSearchRequest,
    LitCompareRequest,
    GapAnalysisRequest,
)
from schemas.ai import EnhanceLiteratureRequest
from utils.auth import get_optional_admin, check_permission
from database import get_db
from models import AdminUser
from api.dependencies import verify_service_access
from services import ai_service, literature_service

lit_router = APIRouter(prefix="/api/literature", tags=["literature"])
lit_compare_router = APIRouter(prefix="/api/lit-compare", tags=["literature-compare"])
enhance_router = APIRouter(prefix="/api", tags=["enhance-literature"])


async def require_literature_access(db: Session, current_user: Optional[AdminUser], token: Optional[str]):
    if current_user:
        check_permission(current_user, "ai")
        return None
    return await verify_service_access(db, token, "ai")


def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract text from uploaded file (PDF, DOCX, TXT)"""
    filename_lower = filename.lower()
    if filename_lower.endswith('.pdf'):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            pages_text = []
            for page in reader.pages[:30]:  # First 30 pages
                pages_text.append(page.extract_text() or "")
            return "\n".join(pages_text)
        except Exception as e:
            return f"[PDF extraction failed: {str(e)}]"
    elif filename_lower.endswith(('.docx', '.doc')):
        try:
            if filename_lower.endswith('.docx'):
                import docx
                doc = docx.Document(io.BytesIO(content))
                return "\n".join([p.text for p in doc.paragraphs])
            else:
                return "[.doc format not supported, please convert to .docx]"
        except Exception as e:
            return f"[DOCX extraction failed: {str(e)}]"
    else:
        return content.decode('utf-8', errors='ignore')[:50000]


@lit_router.post("/search")
async def search_literature_api(request: LiteratureSearchRequest, db: Session = Depends(get_db), current_user: Optional[AdminUser] = Depends(get_optional_admin)):
    """
    Search for real academic literature from multiple databases.

    This endpoint queries PubMed, Europe PMC, CrossRef, and arXiv to find
    real, verifiable academic publications. Results are cached for 1 hour.

    Args:
        request: LiteratureSearchRequest with query, max_results, and databases

    Returns:
        List of citation dictionaries with full metadata
    """
    await require_literature_access(db, current_user, request.token)

    citations = await literature_service.search_literature(
        query=request.query,
        max_results=request.max_results,
        databases=request.databases
    )

    return {
        "query": request.query,
        "count": len(citations),
        "citations": citations
    }


@lit_router.get("/search")
async def search_literature_get(query: str, max_results: int = 10, databases: str = "pubmed,europepmc,crossref", token: Optional[str] = None, db: Session = Depends(get_db)):
    """
    GET endpoint for literature search (for easier browser testing).
    """
    await verify_service_access(db, token, "ai")
    db_list = [d.strip() for d in databases.split(",")]
    citations = await literature_service.search_literature(
        query=query,
        max_results=max_results,
        databases=db_list
    )

    return {
        "query": query,
        "count": len(citations),
        "citations": citations
    }


@enhance_router.post("/enhance-literature-upload")
async def enhance_literature(
    file: Optional[UploadFile] = None,
    text: Optional[str] = Form(None),
    citations: Optional[str] = Form(None),
    token: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """
    Enhance a given paragraph by incorporating citations from uploaded files or provided citation strings.
    Uses multipart/form-data for file uploads.
    """
    if not text:
        raise HTTPException(status_code=400, detail="Missing text parameter")

    await require_literature_access(db, current_user, token)

    citation_content = ""

    if file:
        # Extract text from uploaded file
        content = await file.read()
        filename = file.filename or ""

        if filename.lower().endswith('.pdf'):
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(content))
                pages_text = []
                for page in reader.pages[:10]:  # First 10 pages
                    pages_text.append(page.extract_text() or "")
                citation_content = "\n".join(pages_text)
            except Exception as e:
                citation_content = f"[PDF extraction failed: {str(e)}]"
        elif filename.lower().endswith(('.docx', '.doc')):
            try:
                if filename.lower().endswith('.docx'):
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    citation_content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    citation_content = "[.doc format not supported, please convert to .docx]"
            except Exception as e:
                citation_content = f"[DOCX extraction failed: {str(e)}]"
        else:
            citation_content = content.decode('utf-8', errors='ignore')[:5000]
    elif citations:
        citation_content = citations

    # Build prompt for AI to enhance the text with citations
    prompt = f"""Please enhance the following academic paragraph by incorporating relevant citations and references.

Original text:
{text}

Reference material:
{citation_content}

Requirements:
1. Keep the original meaning and structure
2. Add appropriate in-text citations where claims are made
3. Integrate references naturally into the academic writing style
4. Use standard academic citation format (Author, Year)
5. Maintain formal academic tone
6. Add a References section at the end if new citations were added

Enhanced text:"""

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    enhanced_text = ai_response.get("content", "")

    return {
        "original_text": text,
        "enhanced_text": enhanced_text,
        "citations_used": bool(citation_content)
    }


@enhance_router.post("/enhance-literature")
async def enhance_literature_json(
    req: EnhanceLiteratureRequest,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """
    JSON version of enhance-literature endpoint.
    Expects: {"text": "...", "citations": "..."}
    """
    text = req.text
    citations = req.citations

    if not text:
        raise HTTPException(status_code=400, detail="Missing text parameter")

    await require_literature_access(db, current_user, req.token)

    # Build prompt for AI to enhance the text with citations
    prompt = f"""Please enhance the following academic paragraph by incorporating relevant citations and references.

Original text:
{text}

Reference material:
{citations}

Requirements:
1. Keep the original meaning and structure
2. Add appropriate in-text citations where claims are made
3. Integrate references naturally into the academic writing style
4. Use standard academic citation format (Author, Year)
5. Maintain formal academic tone
6. Add a References section at the end if new citations were added

Enhanced text:"""

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    enhanced_text = ai_response.get("content", "")

    return {
        "original_text": text,
        "enhanced_text": enhanced_text,
        "citations_used": bool(citations)
    }


@lit_compare_router.post("/extract")
async def lit_compare_extract(
    files: List[UploadFile] = File(...),
    token: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """Extract text from uploaded literature files."""
    await require_literature_access(db, current_user, token)

    if len(files) < 1 or len(files) > 5:
        raise HTTPException(status_code=400, detail="请上传 1-5 个文件")

    documents = []
    for file in files:
        content = await file.read()
        text = extract_text_from_file(content, file.filename or "")
        documents.append({
            "name": file.filename or "Unknown",
            "content": text[:30000],  # Limit to 30k chars
        })

    return {"documents": documents}


@lit_compare_router.post("/analyze")
async def lit_compare_analyze(
    req: LitCompareRequest,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """Compare multiple literature documents and provide analysis."""
    await require_literature_access(db, current_user, req.token)

    if len(req.documents) < 2:
        raise HTTPException(status_code=400, detail="Need at least 2 documents")

    # Build comparison prompt
    doc_summaries = []
    for i, doc in enumerate(req.documents):
        content_preview = doc.get("content", "")[:3000]
        doc_summaries.append(f"""
文献 {i+1}: {doc.get("name", "Unknown")}
内容预览:
{content_preview}
""")

    prompt = f"""你是一位资深的学术研究者，请对以下 {len(req.documents)} 篇文献进行详细对比分析。

{chr(10).join(doc_summaries)}

请从以下维度进行对比分析：

## 一、研究主题与目标对比
- 各文献的核心研究问题
- 研究目标的差异

## 二、方法论对比
- 研究设计（实验、调查、综述等）
- 数据采集方法
- 分析技术

## 三、研究发现与结论对比
- 主要发现
- 结论的异同
- 创新点

## 四、文献质量与影响力评估
- 研究方法严谨性
- 数据充分性
- 结论可靠性

## 五、各文献的优劣势分析
- 优势
- 局限性

【重要要求】
- 直接开始输出分析报告内容，严禁任何问候语、客套话或对话式开头（如"好的"、"收到"、"作为一名..."等）。
- 输出必须是纯粹的学术分析文本。
"""

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    comparison = ai_response.get("content", "")

    return {"comparison": comparison}


@lit_compare_router.post("/gap-analysis")
async def lit_gap_analysis(
    req: GapAnalysisRequest,
    db: Session = Depends(get_db),
    current_user: Optional[AdminUser] = Depends(get_optional_admin)
):
    """Analyze gap between start and end documents, provide 4-stage action plan."""
    await require_literature_access(db, current_user, req.token)

    start_preview = req.startDocument.get("content", "")[:3000]
    end_preview = req.endDocument.get("content", "")[:3000]

    comparison_context = req.comparison[:2000] if req.comparison else ""
    comp_section = f"## 文献对比分析参考\n{comparison_context}" if comparison_context else ""

    prompt = f"""你是一位经验丰富的学术研究导师。用户希望从当前水平（起点文献）提升到目标水平（终点文献），请帮助分析差距并制定详细的行动计划。

## 起点文献（当前水平）
标题: {req.startDocument.get("name", "Unknown")}
内容预览:
{start_preview}

## 终点文献（目标水平）
标题: {req.endDocument.get("name", "Unknown")}
内容预览:
{end_preview}

{comp_section}

请进行以下详细分析：

# 一、差距诊断
对比起点与终点文献，识别以下差距：
- 知识储备差距
- 实验/技术能力差距
- 数据分析能力差距
- 写作表达差距
- 研究设计差距

# 二、四阶段行动计划

将差距弥补过程分解为4个阶段（或更细致的子阶段），每个阶段明确：

## 阶段一：基础夯实（预计X周）
- **要做什么**：具体任务清单
- **怎么做**：详细方法和步骤
- **需要什么资源**：文献、工具、软件、课程等
- **如何解决问题**：遇到困难时的解决途径

## 阶段二：技能提升（预计X周）
- **要做什么**
- **怎么做**
- **需要什么资源**
- **如何解决问题**

## 阶段三：实践应用（预计X周）
- **要做什么**
- **怎么做**
- **需要什么资源**
- **如何解决问题**

## 阶段四：成果产出（预计X周）
- **要做什么**
- **怎么做**
- **需要什么资源**
- **如何解决问题**

# 三、整体时间表
| 阶段 | 时间 | 里程碑 |
|------|------|--------|
| ... | ... | ... |

# 四、资源建议
- 必读文献/书籍
- 推荐课程/教程
- 必备软件/工具
- 其他资源

# 五、行动清单（Checklist）
为小白用户提供清晰的每一步行动指引，包括：
- [ ] 每日/每周任务
- [ ] 关键检查点
- [ ] 自我评估方法

【重要要求】
- 直接输出分析结果，严禁任何问候语、客套话或对话式开头（如"好的"、"收到"、"作为您的导师..."等）。
- 内容必须具体、实用，直接以标题开始。
"""

    ai_response = await ai_service.chat_completion([{"role": "user", "content": prompt}])
    analysis = ai_response.get("content", "")

    return {"analysis": analysis}
